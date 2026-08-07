"""Extract raw text from resume files (PDF/DOCX), with an OCR fallback for
scanned/image-only PDFs.

OCR deps (pytesseract, pdf2image) are optional (`pip install resume-ai[ocr]`)
and imported lazily so importing this module never requires the system
`tesseract`/`poppler` binaries unless OCR is actually triggered.
"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any

import pdfplumber
from docx import Document

# Below this many non-whitespace characters, a "text" extraction of a PDF is
# treated as empty (e.g. a scanned page with only a header/footer as real
# text) and OCR fallback is attempted.
_MIN_TEXT_CHARS = 20


class UnsupportedFileTypeError(ValueError):
    pass


@dataclass
class ExtractionResult:
    text: str
    used_ocr: bool = False
    page_count: int | None = None


def extract_text(file_path: str | Path, *, ocr_fallback: bool = True) -> ExtractionResult:
    """Extract text from a resume file. Dispatches on file extension."""
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(path, ocr_fallback=ocr_fallback)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".txt":
        return ExtractionResult(text=path.read_text(encoding="utf-8"))

    raise UnsupportedFileTypeError(f"Unsupported resume file type: {suffix!r}")


def _extract_pdf(path: Path, *, ocr_fallback: bool) -> ExtractionResult:
    with pdfplumber.open(path) as pdf:
        page_count = len(pdf.pages)
        pages_text = [page.extract_text() or "" for page in pdf.pages]

    text = "\n".join(pages_text)
    if len(text.strip()) >= _MIN_TEXT_CHARS or not ocr_fallback:
        return ExtractionResult(text=text, used_ocr=False, page_count=page_count)

    ocr_text = _ocr_pdf(path)
    return ExtractionResult(text=ocr_text, used_ocr=True, page_count=page_count)


def _extract_docx(path: Path) -> ExtractionResult:
    document = Document(str(path))
    paragraphs = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            paragraphs.extend(cell.text for cell in row.cells)
    return ExtractionResult(text="\n".join(paragraphs))


def _pdf_pages_to_images(path: Path) -> list[Any]:
    """Lazy wrapper around pdf2image so the OCR extra is only required here."""
    from pdf2image import convert_from_path

    return convert_from_path(str(path))


def _ocr_image(image: Any) -> str:
    """Lazy wrapper around pytesseract so the OCR extra is only required here."""
    import pytesseract

    return pytesseract.image_to_string(image)


def _ocr_pdf(path: Path) -> str:
    try:
        images = _pdf_pages_to_images(path)
    except ImportError as exc:
        raise RuntimeError(
            "OCR fallback requires the 'ocr' extra (pytesseract, pdf2image) "
            "plus system binaries tesseract and poppler. "
            "Install with: pip install resume-ai[ocr]"
        ) from exc

    return "\n".join(_ocr_image(image) for image in images)
