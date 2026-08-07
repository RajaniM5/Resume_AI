from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from ml.parsing import text_extraction
from ml.parsing.text_extraction import UnsupportedFileTypeError, extract_text


def test_extract_text_from_docx(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Senior Engineer")
    docx_path = tmp_path / "resume.docx"
    doc.save(docx_path)

    result = extract_text(docx_path)

    assert "Jane Doe" in result.text
    assert "Senior Engineer" in result.text
    assert result.used_ocr is False


def test_extract_text_from_txt(tmp_path: Path) -> None:
    txt_path = tmp_path / "resume.txt"
    txt_path.write_text("Plain text resume", encoding="utf-8")

    result = extract_text(txt_path)

    assert result.text == "Plain text resume"


def test_unsupported_extension_raises(tmp_path: Path) -> None:
    bad_path = tmp_path / "resume.exe"
    bad_path.write_bytes(b"not a resume")

    with pytest.raises(UnsupportedFileTypeError):
        extract_text(bad_path)


def _fake_pdf(pages_text: list[str]) -> MagicMock:
    pdf = MagicMock()
    pdf.__enter__.return_value = pdf
    pdf.__exit__.return_value = False
    pdf.pages = []
    for text in pages_text:
        page = MagicMock()
        page.extract_text.return_value = text
        pdf.pages.append(page)
    return pdf


def test_pdf_with_extractable_text_skips_ocr(tmp_path: Path) -> None:
    pdf_path = tmp_path / "resume.pdf"
    pdf_path.touch()
    fake_pdf = _fake_pdf(["This is a resume with plenty of extractable text content."])

    with patch.object(text_extraction.pdfplumber, "open", return_value=fake_pdf):
        result = extract_text(pdf_path)

    assert "extractable text" in result.text
    assert result.used_ocr is False
    assert result.page_count == 1


def test_pdf_with_empty_text_falls_back_to_ocr(tmp_path: Path) -> None:
    pdf_path = tmp_path / "scanned.pdf"
    pdf_path.touch()
    fake_pdf = _fake_pdf([""])

    with (
        patch.object(text_extraction.pdfplumber, "open", return_value=fake_pdf),
        patch.object(text_extraction, "_pdf_pages_to_images", return_value=["page-1-image"]),
        patch.object(text_extraction, "_ocr_image", return_value="OCR'd resume text"),
    ):
        result = extract_text(pdf_path)

    assert result.text == "OCR'd resume text"
    assert result.used_ocr is True


def test_pdf_with_empty_text_and_ocr_disabled_returns_empty(tmp_path: Path) -> None:
    pdf_path = tmp_path / "scanned.pdf"
    pdf_path.touch()
    fake_pdf = _fake_pdf([""])

    with patch.object(text_extraction.pdfplumber, "open", return_value=fake_pdf):
        result = extract_text(pdf_path, ocr_fallback=False)

    assert result.text == ""
    assert result.used_ocr is False


def test_ocr_missing_deps_raises_clear_error(tmp_path: Path) -> None:
    pdf_path = tmp_path / "scanned.pdf"
    pdf_path.touch()
    fake_pdf = _fake_pdf([""])

    with (
        patch.object(text_extraction.pdfplumber, "open", return_value=fake_pdf),
        patch.object(text_extraction, "_pdf_pages_to_images", side_effect=ImportError()),
    ):
        with pytest.raises(RuntimeError, match="ocr"):
            extract_text(pdf_path)
